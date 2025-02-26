import streamlit as st
import os
import re
import json
import zipfile
import shutil
import tempfile
from PIL import Image
from docx import Document
import fitz  # PyMuPDF
import requests

# ---- Pour l'extraction de texte depuis le PDF (langchain_community) ----
# Veillez à avoir installé la dépendance : pip install langchain_community
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter

# =======================================================================
# Configuration Azure OpenAI (clé API et URL du endpoint)
# =======================================================================

API_KEY = st.secrets["Clé secrète"]
AZURE_ENDPOINT = st.secrets["Lien connexion"]
DEPLOYMENT_NAME = st.secrets["gpt-4o-mini"]

def azure_llm_predict(prompt):
    headers = {
        "Content-Type": "application/json",
        "api-key": API_KEY
    }
    messages = [
        {"role": "system", "content": "Tu es un assistant utile."},
        {"role": "user", "content": prompt}
    ]
    data = {
        "messages": messages,
        "max_tokens": 2048
    }
    response = requests.post(AZURE_ENDPOINT, headers=headers, json=data)
    if response.status_code == 200:
        result = response.json()
        return result["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Erreur {response.status_code}: {response.text}")

def extract_text_from_pdf(pdf_path):
    pdf_loader = PyPDFLoader(pdf_path)
    data = pdf_loader.load()
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=0)
    return text_splitter.split_documents(data)

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip() != "")
    class SimpleDoc:
        def __init__(self, text):
            self.page_content = text
    splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=0)
    splits = splitter.split_text(full_text)
    return [SimpleDoc(t) for t in splits]

def extract_information(splits):
    cv_text = "\n".join([split.page_content for split in splits])
    template = """
    Vous allez extraire les informations du CV ci-dessous et retourner UNIQUEMENT du JSON,
    sans aucune explication ni commentaire.

    ### Structure du JSON à suivre (pas d'écart autorisé) : ###
    {{
      "Prénom": "",
      "Nom": "",
      "Poste convoité": "",
      "Années d’expérience": "",
      "Téléphone": "",
      "Adresse": "",
      "Email": "",
      "LinkedIn": "",
      "Description de la personne": "",
      "Compétences techniques": [],
      "Technologies": [],
      "Langues": {{ }},
      "Expérience professionnelle": [
        {{
          "Entreprise": "",
          "Poste": "",
          "Durée": "",
          "Contexte": "",
          "Missions": [],
          "Environnement": []
        }}
      ],
      "Formations": [
        {{
          "Nom": "",
          "Institution": "",
          "Durée": ""
        }}
      ]
    }}

    ### Exemples : ###
    # EXEMPLE 1 :
    {{
      "Prénom": "Alice",
      "Nom": "Durand",
      "Poste convoité": "Consultant CRM",
      "Années d’expérience": "5",
      "Téléphone": "+33 6 12 34 56 78",
      "Adresse": "45 Rue de la République, 75010 Paris",
      "Email": "alice.durand@example.com",
      "LinkedIn": "https://www.linkedin.com/in/alice-durand/",
      "Description de la personne": "Spécialiste CRM et marketing automation, forte appétence data.",
      "Compétences techniques": [
        "SQL",
        "Python",
        "Marketing Automation",
        "Salesforce",
        "HubSpot"
      ],
      "Technologies": [
        "AWS",
        "GCP",
        "Power BI"
      ],
      "Langues": {{
        "Français": "Bilingue",
        "Anglais": "Courant"
      }},
      "Expérience professionnelle": [
        {{
          "Entreprise": "Klint",
          "Poste": "Consultant CRM",
          "Durée": "Janvier 2021 - Présent",
          "Contexte": "Projet de mise en place d'un nouveau CRM marketing pour client retail",
          "Missions": [
            "Analyser les besoins métiers",
            "Intégration du CRM HubSpot et migration des données",
            "Formation des équipes internes"
          ],
          "Environnement": [
            "HubSpot",
            "AWS",
            "SQL",
            "Looker"
          ]
        }}
      ],
      "Formations": [
        {{
          "Nom": "Master en Marketing Digital",
          "Institution": "Université de Lyon",
          "Durée": "2015-2017"
        }}
      ]
    }}

    # EXEMPLE 2 :
    {{
      "Prénom": "Jean",
      "Nom": "Moreau",
      "Poste convoité": "Data Engineer",
      "Années d’expérience": "7",
      "Téléphone": "+33 7 54 32 11 22",
      "Adresse": "12 Avenue Data, 69003 Lyon",
      "Email": "jean.moreau@example.com",
      "LinkedIn": "https://www.linkedin.com/in/jean-moreau/",
      "Description de la personne": "Expert en pipelines et intégration big data, solutions cloud.",
      "Compétences techniques": [
        "Python",
        "Spark",
        "ETL",
        "Kafka"
      ],
      "Technologies": [
        "Azure",
        "Databricks",
        "Airflow",
        "Power BI"
      ],
      "Langues": {{
        "Français": "Bilingue",
        "Anglais": "TOEIC 920"
      }},
      "Expérience professionnelle": [
        {{
          "Entreprise": "Grande Banque",
          "Poste": "Data Engineer",
          "Durée": "2019 - 2023",
          "Contexte": "Implémentation d'un data lake sur Azure pour la direction marketing.",
          "Missions": [
            "Mise en place d'un pipeline d'intégration temps réel",
            "Optimisation de jobs Spark",
            "Développement de dashboards Power BI"
          ],
          "Environnement": [
            "Azure",
            "Databricks",
            "Kafka",
            "Spark"
          ]
        }}
      ],
      "Formations": [
        {{
          "Nom": "Master Informatique - Big Data",
          "Institution": "Université Paris-Saclay",
          "Durée": "2012-2014"
        }}
      ]
    }}

    ### CONSIGNES SPÉCIFIQUES :
    - Remplir "Compétences techniques" avec les langages, frameworks, etc.
    - Remplir "Technologies" avec les plateformes ou solutions globales.
    - Pour "Missions", EXTRAIRE TOUTES les missions telles qu'elles figurent dans le CV, sans les tronquer ni les résumer. Veuillez également inclure toutes les missions des sections additionnelles (ex. "Projet 1", "Projet 2", etc.).
    - Remplir "Contexte" avec quelques mots décrivant la situation globale du projet.
    - Classer les expériences de la plus récente à la plus ancienne.

    ### À VOUS DE JOUER ###
    Analysez le CV ci-dessous et retournez UNIQUEMENT le JSON,
    sans aucune explication ou commentaire supplémentaire.
    {cv_text}
    """
    prompt = template.format(cv_text=cv_text)
    return azure_llm_predict(prompt)

def is_json_complete(json_str):
    return json_str.count("{") == json_str.count("}")

def complete_json(partial_json):
    continuation_prompt = (
        "Le JSON précédent est incomplet. Merci de continuer le JSON en reprenant exactement là où il s'est arrêté, "
        "sans ajouter de texte supplémentaire, afin de fournir l'intégralité des informations demandées.\n\n"
        "JSON partiel :\n" + partial_json
    )
    return azure_llm_predict(continuation_prompt)

def get_complete_json(initial_response, max_attempts=3):
    complete_response = initial_response
    attempts = 0
    while not is_json_complete(complete_response) and attempts < max_attempts:
        extra = complete_json(complete_response)
        complete_response += extra
        attempts += 1
    return complete_response

def parse_result(result):
    match = re.search(r'(\{.*\})', result, flags=re.DOTALL)
    if match:
        json_str = match.group(1)
        if not is_json_complete(json_str):
            json_str = get_complete_json(json_str)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            raise Exception(f"Erreur de décodage JSON: {e}")
    else:
        raise Exception("Aucun JSON détecté dans la réponse du LLM.")

def parse_start_date(duration_str):
    try:
        start_part = duration_str.split("-")[0].strip()
        parts = start_part.split()
        if len(parts) >= 2:
            month_abbr = parts[0]
            year = int(parts[1])
            month_map = {
                "Jan": 1, "Feb": 2, "Mar": 3, "Apr": 4,
                "May": 5, "Jun": 6, "Jul": 7, "Aug": 8,
                "Sep": 9, "Oct": 10, "Nov": 11, "Dec": 12
            }
            month = month_map.get(month_abbr, 0)
            return (year, month)
    except Exception as e:
        raise Exception(f"Erreur lors du parsing de la durée '{duration_str}': {e}")
    return (0, 0)

def sort_experiences(exp_list):
    return sorted(exp_list, key=lambda exp: parse_start_date(exp.get("Durée", "")), reverse=True)

def replace_placeholder_in_runs(paragraph, placeholder, new_text):
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False
    replaced_text = full_text.replace(placeholder, new_text)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced_text
    else:
        paragraph.add_run(replaced_text)
    return True

def replace_in_paragraphs_and_tables_runs(doc_obj, placeholder, value):
    replaced_something = False
    for paragraph in doc_obj.paragraphs:
        if replace_placeholder_in_runs(paragraph, placeholder, value):
            replaced_something = True
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_placeholder_in_runs(paragraph, placeholder, value):
                        replaced_something = True
    return replaced_something

def remove_unreplaced_placeholders(doc):
    placeholder_pattern = re.compile(r'{{.*?}}')
    environment_pattern = re.compile(r'{{E\s+\d+\s+E\s+\d+}}')

    def paragraph_contains_non_env_placeholder(paragraph_text):
        all_placeholders = placeholder_pattern.findall(paragraph_text)
        non_env = [ph for ph in all_placeholders if not environment_pattern.fullmatch(ph)]
        return len(non_env) > 0

    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if paragraph_contains_non_env_placeholder(paragraph.text):
            paragraphs_to_remove.append(paragraph)

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs_in_cell = list(cell.paragraphs)
                for paragraph in paragraphs_in_cell:
                    if paragraph_contains_non_env_placeholder(paragraph.text):
                        paragraph._element.getparent().remove(paragraph._element)

def remove_unused_environment_placeholders(doc):
    environment_pattern = re.compile(r'{{E\s+\d+\s+E\s+\d+}}')
    for paragraph in doc.paragraphs:
        if environment_pattern.search(paragraph.text):
            paragraph.text = environment_pattern.sub('', paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if environment_pattern.search(paragraph.text):
                        paragraph.text = environment_pattern.sub('', paragraph.text)

def fill_klint_cv_template(response_data, template_path, output_path):
    doc = Document(template_path)
    cv_info = {
        'Prénom': response_data.get('Prénom', ''),
        'Nom': response_data.get('Nom', ''),
        'Poste convoité': response_data.get('Poste convoité', ''),
        'Années d’expérience': response_data.get('Années d’expérience', ''),
        'Téléphone': response_data.get('Téléphone', ''),
        'Adresse': response_data.get('Adresse', ''),
        'Email': response_data.get('Email', ''),
        'LinkedIn': response_data.get('LinkedIn', ''),
        'Description de la personne': response_data.get('Description de la personne', ''),
        'Compétences techniques': response_data.get('Compétences techniques', []),
        'Technologies': response_data.get('Technologies', []),
        'Langues': response_data.get('Langues', {}),
        'Expérience professionnelle': response_data.get('Expérience professionnelle', []),
        'Formations': response_data.get('Formations', [])
    }

    cv_info['Expérience professionnelle'] = sort_experiences(cv_info.get('Expérience professionnelle', []))

    def replace_placeholder(placeholder, value):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

    placeholders_map = {
        '{{Prénom}}': cv_info['Prénom'],
        '{{Nom}}': cv_info['Nom'],
        '{{Poste convoité}}': cv_info['Poste convoité'],
        '{{Année d’expérience}}': f"{cv_info['Années d’expérience']} années d'expérience",
        '{{Téléphone}}': cv_info['Téléphone'],
        '{{Adresse}}': cv_info['Adresse'],
        '{{Email}}': cv_info['Email'],
        '{{LinkedIn}}': cv_info['LinkedIn'],
        '{{Description de la personne}}': cv_info['Description de la personne']
    }

    for ph, val in placeholders_map.items():
        replace_placeholder(ph, val)

    replace_placeholder('{{COMPETENCES}}', '\n'.join(cv_info['Compétences techniques']))
    replace_placeholder('{{TECHNOLOGIES}}', ' / '.join(cv_info['Technologies']))

    langues_str = '\n'.join([f"{lang}: {niveau}" for lang, niveau in cv_info['Langues'].items()])
    replace_placeholder('{{LANGUES}}', langues_str)

    for exp_idx, exp in enumerate(cv_info['Expérience professionnelle'], start=1):
        replace_placeholder(f'{{{{Poste expérience {exp_idx}}}}}', exp.get('Poste', ''))
        replace_placeholder(f'{{{{Entreprise expérience {exp_idx}}}}}', exp.get('Entreprise', ''))
        replace_placeholder(f'{{{{Date expérience {exp_idx}}}}}', exp.get('Durée', ''))
        replace_placeholder(f'{{{{Contexte {exp_idx}}}}}', exp.get('Contexte', ''))

        missions_list = exp.get('Missions', [])
        missions_str = "\n".join("- " + mission for mission in missions_list)
        replace_placeholder(f'{{{{MISSIONS {exp_idx}}}}}', missions_str)

        environnements_str = ' - '.join(exp.get('Environnement', []))
        replace_placeholder(f'{{{{ENVIRONNEMENT {exp_idx}}}}}', environnements_str)

    for f_idx, formation in enumerate(cv_info['Formations'], start=1):
        replace_placeholder(f'{{{{ANNEE D’OBTENTION {f_idx}}}}}', formation.get('Durée', ''))
        formation_str = f"{formation.get('Nom', '')} - {formation.get('Institution', '')}"
        replace_placeholder(f'{{{{Intitulé du diplôme {f_idx}}}}}', formation_str)

    remove_unreplaced_placeholders(doc)
    doc.save(output_path)

def extract_first_image_from_pdf(pdf_path, output_image_path):
    doc_pdf = fitz.open(pdf_path)
    try:
        for page_index in range(len(doc_pdf)):
            page = doc_pdf[page_index]
            images_info = page.get_images(full=True)
            for img_info in images_info:
                xref = img_info[0]
                base_image = doc_pdf.extract_image(xref)
                image_data = base_image["image"]
                with open(output_image_path, "wb") as f:
                    f.write(image_data)
                return True
        return False
    finally:
        doc_pdf.close()

def resize_image(input_image_path, output_image_path, size):
    with Image.open(input_image_path) as img:
        img = img.resize((size, size), Image.Resampling.LANCZOS)
        img.save(output_image_path, format="PNG")

def unzip_docx(docx_path, extract_folder):
    if os.path.exists(extract_folder):
        shutil.rmtree(extract_folder)
    os.makedirs(extract_folder, exist_ok=True)
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(extract_folder)

def rezip_docx(folder_to_zip, output_docx):
    with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as new_zip:
        for folder_name, subfolders, filenames in os.walk(folder_to_zip):
            for filename in filenames:
                file_path = os.path.join(folder_name, filename)
                arcname = os.path.relpath(file_path, folder_to_zip)
                new_zip.write(file_path, arcname)

def replace_shape_image_in_docx(docx_path, new_image_path, docx_output):
    temp_folder = "temp_docx_unzipped"
    unzip_docx(docx_path, temp_folder)
    media_folder = os.path.join(temp_folder, "word", "media")
    if not os.path.exists(media_folder):
        return
    target_image = "image3.png"
    old_image_path = os.path.join(media_folder, target_image)
    if os.path.exists(old_image_path):
        resized_image_path = "resized_image.png"
        resize_image(new_image_path, resized_image_path, size=500)
        shutil.copy(resized_image_path, old_image_path)
        if os.path.exists(resized_image_path):
            os.remove(resized_image_path)
    rezip_docx(temp_folder, docx_output)
    shutil.rmtree(temp_folder, ignore_errors=True)

def insert_photo_into_cv(docx_path, pdf_path, output_docx):
    extracted_image = "extracted_photo.png"
    if not extract_first_image_from_pdf(pdf_path, extracted_image):
        return False
    replace_shape_image_in_docx(docx_path, extracted_image, output_docx)
    if os.path.exists(extracted_image):
        os.remove(extracted_image)
    return True

def page_formatage_cv():
    """
    Cette fonction gère la page de Service 1 (Formatage de CV)
    où on applique tout l'algo.
    """
    st.markdown("""
    <style>
        .title {
            font-size: 2rem;
            font-weight: bold;
            color: #007ACC;
            text-align: center;
            padding-bottom: 50px;
        }
        .section-header {
            font-size: 1rem;
            font-weight: bold;
            color: #333333;
            border-bottom: 2px solid #007ACC;
            padding-bottom: 5px;
        }
        .container {
            padding: 20px;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div class='title'>CV Brut - CV Klint</div>", unsafe_allow_html=True)

    col_import, col_download = st.columns(2)

    with col_import:
        st.markdown("<div class='section-header'>Importer votre CV</div>", unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Choisissez votre CV (PDF ou WORD)", type=["pdf", "docx"])

    # On n'affiche plus de champ texte : on va fixer le nom du fichier final

    # Pour conserver les bytes du fichier final en mémoire (session) :
    if "final_doc_bytes" not in st.session_state:
        st.session_state["final_doc_bytes"] = None

    # Nouveau bouton : on ne lance l'algorithme qu'au clic
    if uploaded_file is not None:
        temp_dir = tempfile.mkdtemp()
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        input_path = os.path.join(temp_dir, "input" + file_ext)
        with open(input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Bouton pour lancer le traitement
        if st.button("2) Lancer le formatage du CV"):
            with st.spinner("Traitement en cours..."):
                # Extraction du texte
                if file_ext == ".pdf":
                    splits = extract_text_from_pdf(input_path)
                elif file_ext == ".docx":
                    splits = extract_text_from_docx(input_path)
                else:
                    st.error("Format de fichier non supporté.")
                    return

                # Extraction des infos via LLM
                result_str = extract_information(splits)
                response_json = parse_result(result_str)

                # Sauvegarde du JSON (optionnel)
                json_output_path = os.path.join(temp_dir, "cv_extrait.json")
                with open(json_output_path, "w", encoding="utf-8") as f:
                    json.dump(response_json, f, ensure_ascii=False, indent=4)

                # Chemin vers le template Word
                template_path = "Modèle_klint_remplir3.docx"
                if not os.path.exists(template_path):
                    st.error("Le template Word 'Modèle_klint_remplir3.docx' est introuvable.")
                    return

                # Création du CV final
                first_name = response_json.get("Prénom", "").strip()
                last_name = response_json.get("Nom", "").strip()
                first_name_safe = re.sub(r"[^\w\- ]", "_", first_name)
                last_name_safe = re.sub(r"[^\w\- ]", "_", last_name)
                final_output_path = os.path.join(temp_dir, f"cv_{first_name_safe}_{last_name_safe}.docx")

                # Remplissage du template
                fill_klint_cv_template(response_json, template_path, final_output_path)

                # Ajout automatique de la photo si le CV d'origine en contenait une (PDF)
                insert_photo_into_cv(final_output_path, input_path, final_output_path)

                # On lit le fichier final en bytes
                with open(final_output_path, "rb") as final_file:
                    st.session_state["final_doc_bytes"] = final_file.read()

                st.success("CV formaté avec succès ! Vous pouvez maintenant le télécharger.")

    with col_download:
        st.markdown("<div class='section-header'>Télécharger le CV Klint</div>", unsafe_allow_html=True)

        # On n'affiche le bouton de téléchargement que s'il y a un fichier généré
        if st.session_state["final_doc_bytes"] is not None:
            st.download_button(
                "Télécharger le CV",
                data=st.session_state["final_doc_bytes"],
                file_name="cv_formate.docx",  # Nom fixé
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.info("Aucun CV formaté disponible pour le moment. Veuillez d'abord importer un fichier et lancer le traitement.")

def azure_llm_chat(prompt):
    """
    Envoie une requête à Azure OpenAI et retourne la réponse.
    """
    headers = {
        "Content-Type": "application/json",
        "api-key": API_KEY
    }
    messages = [
        {"role": "system", "content": "Tu es un assistant IA utile."},
        {"role": "user", "content": prompt}
    ]
    data = {
        "messages": messages,
        "max_tokens": 2048
    }
    
    response = requests.post(AZURE_ENDPOINT, headers=headers, json=data)
    
    if response.status_code == 200:
        result = response.json()
        return result["choices"][0]["message"]["content"]
    else:
        return f"Erreur {response.status_code}: {response.text}"

def page_chat():
    """
    Interface pour le service de Chat.
    """
    st.title("💬 Chat avec Azure OpenAI")

    st.write("Posez une question et obtenez une réponse générée par Azure OpenAI.")

    user_input = st.text_area("Entrez votre message :", "")

    if st.button("Envoyer"):
        if user_input:
            with st.spinner("Azure OpenAI réfléchit..."):
                response = azure_llm_chat(user_input)
            st.subheader("Réponse :")
            st.write(response)
        else:
            st.warning("Veuillez entrer une question.")

def page_service_3():
    """
    Page de Service 3 (vide - en preview).
    """
    st.title("Service 3 - Preview")
    st.write("Cette page est en cours de développement... Patience !")

def main():
    # Barre latérale pour la navigation
    st.sidebar.title("Menu de navigation")
    choix = st.sidebar.radio(
        "Choisissez un service :",
        ("Formatage de CV", "Service 2", "Service 3")
    )

    # Selon le choix, on appelle la page correspondante
    if choix == "Formatage de CV":
        page_formatage_cv()
    elif choix == "Service 2":
        page_chat()
    elif choix == "Service 3":
        page_service_3()

if __name__ == "__main__":
    main()
